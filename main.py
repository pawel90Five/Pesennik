from docx.text.paragraph import Paragraph
from pypdf import PdfReader, PdfWriter, Transformation
from utils import Song, less_spaces, is_header_line

import pymupdf  # PyMuPDF


def save_pages(pdf_file, new_pdf_file, start, to):
    """function that save pages from pdf file to another pdf file
    with same formating"""
    reader = PdfReader(pdf_file)

    writer = PdfWriter()

    for i in range(start, to):
        writer.add_page(reader.pages[i])

    writer.write(open(new_pdf_file, 'wb'))


def collect_songs(pdf_file):
    doc = pdf_file
    songs = {}
    ls = 1
    header = ""
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] == 0:
                
                for line in block["lines"]:
                    sentence = []
                    is_header = False
                    
                    bruh = 0
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if text:
                            if not sentence:
                                bold = "bold" if "bold" in span["font"].lower() else ""
                                sentence.append((span["origin"],
                                                 bold))
                            if is_header_line(line["spans"], 10) or is_header: # не факт
                                is_header = True
                                print(span["text"], end=" ")
                                bruh += 1
                            sentence.append(span["text"])
                            
                    if sentence:
                        if is_header:
                            header = less_spaces(" ".join(sentence[1:]), ls)
                            if header in songs:
                                songs[header + " 2"] = [sentence[0]]
                            else: 
                                songs[header] = [sentence[0]]
                        else:
                            songs[header].append((sentence[0], less_spaces(" ".join(sentence[1:]), ls)))
                        sentence = []
                    if bruh:        
                        print()
    songs_list:list[Song]=[]
    for song in songs:
        songs_list.append(Song(
            header=song.strip(),
            header_coords=songs[song][0][0],
            text=songs[song][1:]
        ))
    return songs_list

def insert_songs(songs:list[Song]):
    xy = (50, 50)
    doc = pymupdf.Document()
    [doc.new_page() for _ in range(len(songs))]
    
    for j, song in enumerate(songs):
        page = doc[j]
        song.insert_to_page(page, xy)

    doc.save("output.pdf")


def scaling(reader: PdfReader, writer: PdfWriter):

    pages = reader.pages[:]
    pages[0].scale_by(0.5)

    op = Transformation().scale(sx=0.7, sy=0.7)
    pages[1].add_transformation(op)

    writer.add_page(pages[0])
    writer.add_page(pages[1])
    writer.add_page(pages[2])
    writer.write("test.pdf")

def get_text(reader: PdfReader):
    return reader.pages[0].extract_text(extraction_mode = "plain")

def main():
    #reader = PdfReader('src/кусок2.pdf')
    reader = PdfReader('src/песенник.pdf')
    #scaling(reader, writer)
    #print(get_text(reader))
    songs = collect_songs(pymupdf.Document('src/песенник.pdf'))


    insert_songs(songs)

  

if __name__ == '__main__':
    # doc = docx.Document('src/кусок.docx')
    # print_paragraphs(doc)
    #save_3_pages('src/песенник.pdf', 'src/кусок2.pdf')

    main()

    

